import tkinter as tk
from tkinter import messagebox, Frame, filedialog
import pandas as pd
import matplotlib.pyplot as plt
from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg
import seaborn as sns
from docx import Document
from docx.shared import Inches

# Function to save the data to CSV
def save_to_csv():
    data = {
        "PatientGender": entry_gender.get(),
        "PatientAge": entry_age.get(),
        "PatientRegion": entry_address.get(),
        "Doctor consulted ": entry_doctor.get(),
        "Department": entry_department.get(),
        "Case description": entry_case.get(),
        "Severity": entry_severity.get(),
        "Bill": entry_bill.get(),
        "Insurance":entry_Insurance.get(),
        "Final bill": entry_bill.get(),
        "In date":entry_date.get(),
        "Out date":entry_date.get(),
    }
    df = pd.DataFrame([data])
    df.to_csv('finaldata.csv', mode='a', index=False, header=False)
    messagebox.showinfo("Success", "Data saved successfully!")

# Example charts
def chart1():
    df = pd.read_csv('finaldata.csv')
    fig, ax = plt.subplots()
    sns.set(style="whitegrid")
    sns.countplot(y='Case description', data=df, palette='viridis', order=df['Case description'].value_counts().index, ax=ax)
    ax.set_title('Most Common Case Descriptions Treated')
    display_chart(fig)

def chart2():
    df = pd.read_csv('finaldata.csv')
    fig, ax = plt.subplots()
    region_counts = df['PatientRegion'].value_counts().reset_index()
    region_counts.columns = ['PatientRegion', 'Count']
    region_counts = region_counts.sort_values('PatientRegion')
    ax.fill_between(region_counts['PatientRegion'], region_counts['Count'], color="skyblue", alpha=0.4)
    ax.plot(region_counts['PatientRegion'], region_counts['Count'], color="Slateblue", alpha=0.6, linewidth=2)
    ax.set_title('Number of Visitors from Each Region')
    display_chart(fig)

def chart3():
    df = pd.read_csv('finaldata.csv')
    fig, ax = plt.subplots()
    df['Age Category'] = pd.cut(df['PatientAge'], bins=[0, 18, 50, 100], labels=['Children', 'Adults', 'Senior Citizens'])
    age_category_counts = df['Age Category'].value_counts()
    ax.pie(age_category_counts, labels=age_category_counts.index, autopct='%1.1f%%', startangle=140, colors=['#ff9999', '#66b3ff', '#99ff99'])
    ax.set_title('Distribution of Patients by Age Category')
    display_chart(fig)

def chart4():
    df = pd.read_csv('finaldata.csv')
    fig, ax = plt.subplots()
    department_counts = df['Department'].value_counts().reset_index()
    department_counts.columns = ['Department', 'Count']
    ax.bar(department_counts['Department'], department_counts['Count'], color='skyblue')
    ax.set_title('Distribution of Patients per Department')
    display_chart(fig)

def chart7():
    df = pd.read_csv('finaldata.csv')
    fig, ax = plt.subplots()
    severity_distribution = df['Severity'].value_counts()
    ax.pie(severity_distribution, labels=severity_distribution.index, autopct='%1.1f%%', startangle=140, colors=sns.color_palette('coolwarm', len(severity_distribution)))
    ax.set_title('Distribution of Severity of Cases')
    display_chart(fig)

def chart8():
    df = pd.read_csv('finaldata.csv')
    fig, ax = plt.subplots()
    insurance_by_age = df.groupby('PatientAge')['Insurance'].mean().reset_index()
    ax.plot(insurance_by_age['PatientAge'], insurance_by_age['Insurance'], marker='o')
    ax.set_title('Average Insurance Distribution by Patient Age')
    display_chart(fig)

def chart9():
    df = pd.read_csv('finaldata.csv')
    fig, ax = plt.subplots()
    expenditure_by_age = df.groupby('PatientRegion')['Final bill'].sum().reset_index()
    ax.plot(expenditure_by_age['PatientRegion'], expenditure_by_age['Final bill'], marker='o', color='b')
    ax.set_title('Total Amount Spent by Patient Age')
    display_chart(fig)

def chart10():
    df = pd.read_csv('finaldata.csv')
    fig, ax = plt.subplots()
    pivot_table = df.pivot_table(index='Doctor consulted', columns='Case description', aggfunc='size', fill_value=0)
    pivot_table.plot(kind='bar', stacked=True, colormap='viridis', figsize=(14, 8), ax=ax)
    ax.set_title('Number of Cases Each Doctor Has Handled for Each Type of Case Description')
    display_chart(fig)

root = tk.Tk()
root.title("Analytics Input Form")

chart_frame = tk.Frame(root)
chart_frame.grid(row=50, column=0, columnspan=2)

# Function to display the chart in the Tkinter window
def display_chart(fig):
    for widget in chart_frame.winfo_children():
        widget.destroy()
    canvas = FigureCanvasTkAgg(fig, master=chart_frame)
    canvas.draw()
    canvas.get_tk_widget().pack()

charts = [chart1, chart2, chart3, chart4, chart7, chart8, chart9, chart10]
chart_index = 0

# Function to display the next chart
def show_next_chart():
    global chart_index
    if chart_index < len(charts):
        charts[chart_index]()
        chart_index += 1
    else:
        messagebox.showinfo("Info", "No more charts to display.")

# Function to generate and save the report with conclusions
def generate_report():
    df = pd.read_csv('finaldata.csv')
    conclusions = []

    # Example conclusions
    case_counts = df['Case description'].value_counts()
    most_common_case = case_counts.idxmax()
    most_common_case_count = case_counts.max()
    conclusions.append(f"The most common case description is '{most_common_case}' with {most_common_case_count} cases.")

    region_counts = df['PatientRegion'].value_counts()
    most_common_region = region_counts.idxmax()
    most_common_region_count = region_counts.max()
    conclusions.append(f"The region with the most patients is '{most_common_region}' with {most_common_region_count} patients.")

    age_distribution = df['PatientAge'].describe()
    conclusions.append(f"The average age of patients is {age_distribution['mean']:.2f} years.")

    department_counts = df['Department'].value_counts()
    most_common_department = department_counts.idxmax()
    most_common_department_count = department_counts.max()
    conclusions.append(f"The department with the most patients is '{most_common_department}' with {most_common_department_count} patients.")

    # Save conclusions in a document
    doc = Document()
    doc.add_heading('Doctor and Disease Relation Analysis', 0)

    doc.add_heading('Conclusions:', level=1)
    for conclusion in conclusions:
        doc.add_paragraph(conclusion)

    # Save the document
    doc_file = 'doctor_disease_analysis.docx'
    doc.save(doc_file)

    messagebox.showinfo("Report Generated", f"Report saved as {doc_file}")

# Creating the input fields
tk.Label(root, text="Name").grid(row=0)
tk.Label(root, text="Age").grid(row=1)
tk.Label(root, text="Gender").grid(row=2)
tk.Label(root, text="Address").grid(row=3)
tk.Label(root, text="Case").grid(row=4)
tk.Label(root, text="Department").grid(row=5)
tk.Label(root, text="Doctor Name").grid(row=6)
tk.Label(root, text="Severity").grid(row=7)
tk.Label(root, text="Bill Amount").grid(row=8)

entry_name = tk.Entry(root)
entry_age = tk.Entry(root)
entry_gender = tk.Entry(root)
entry_address = tk.Entry(root)
entry_case = tk.Entry(root)
entry_department = tk.Entry(root)
entry_doctor = tk.Entry(root)
entry_severity = tk.Entry(root)
entry_bill = tk.Entry(root)

entry_name.grid(row=0, column=1)
entry_age.grid(row=1, column=1)
entry_gender.grid(row=2, column=1)
entry_address.grid(row=3, column=1)
entry_case.grid(row=4, column=1)
entry_department.grid(row=5, column=1)
entry_doctor.grid(row=6, column=1)
entry_severity.grid(row=7, column=1)
entry_bill.grid(row=8, column=1)

# Creating buttons
tk.Button(root, text="Save to CSV", command=save_to_csv).grid(row=9, column=0)
tk.Button(root, text="Next Chart", command=show_next_chart).grid(row=9, column=1)
tk.Button(root, text="Generate Report", command=generate_report).grid(row=10, column=0, columnspan=2)

root.mainloop()